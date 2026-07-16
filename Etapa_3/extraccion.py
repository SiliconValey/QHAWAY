"""Extracción de contenido desde archivos reales (ING-02, ING-04, ING-06).

Adaptadores hacia el contrato `dominio.contenido`:
* PDF  -> PyMuPDF (fitz)
* docx -> python-docx
* .ui  -> ElementTree (XML, biblioteca estándar)

ING-06: un archivo corrupto, protegido o sin texto extraíble se reporta como
problema **sin abortar** el resto de la entrega. Por eso `extraer` nunca levanta
por un archivo malo: devuelve un `ResultadoExtraccion` con `problema` cargado.
ING-02: los formatos no soportados se rechazan con mensaje.

La heurística de encabezados (Etapa 0.2) arma las secciones numeradas que luego
usan DET-03 (secciones numeradas) y las referencias por ubicación.
"""

from __future__ import annotations

import re
import xml.etree.ElementTree as ET
from dataclasses import dataclass
from pathlib import Path

from ..dominio.contenido import ArbolUI, ContenidoDocumento, NodoUI, Seccion

FORMATOS_SOPORTADOS = {".pdf", ".docx", ".ui"}

# Encabezado numerado: "3 Título", "3.1. Título"... título corto, sin punto final.
_RE_ENCABEZADO = re.compile(r"^\s*(\d+(?:\.\d+)*)\.?\s+(\S.{0,80})$")


@dataclass(frozen=True)
class ResultadoExtraccion:
    """Contenido extraído, o el problema encontrado (ING-06). Nunca ambos nulos."""

    ruta: str
    contenido: ContenidoDocumento | ArbolUI | None = None
    problema: str | None = None

    @property
    def ok(self) -> bool:
        return self.problema is None and self.contenido is not None


def extraer(ruta: Path | str, tipo_artefacto: str) -> ResultadoExtraccion:
    """Extrae el contenido de un archivo según su extensión.

    No levanta por archivos problemáticos (ING-06): los reporta en el resultado.
    """
    ruta = Path(ruta)
    ext = ruta.suffix.lower()

    if ext not in FORMATOS_SOPORTADOS:
        return ResultadoExtraccion(
            ruta=str(ruta),
            problema=(
                f"Formato no soportado: '{ext or ruta.name}'. "
                f"Formatos válidos: {', '.join(sorted(FORMATOS_SOPORTADOS))}."
            ),
        )

    try:
        if ext == ".pdf":
            return _extraer_pdf(ruta, tipo_artefacto)
        if ext == ".docx":
            return _extraer_docx(ruta, tipo_artefacto)
        return _extraer_ui(ruta, tipo_artefacto)
    except Exception as exc:  # ING-06: cualquier falla se reporta, no aborta
        return ResultadoExtraccion(
            ruta=str(ruta), problema=f"No se pudo extraer contenido: {exc}"
        )


# ----------------------------------------------------------------------------
# PDF
# ----------------------------------------------------------------------------
def _extraer_pdf(ruta: Path, tipo_artefacto: str) -> ResultadoExtraccion:
    import fitz  # PyMuPDF

    doc = fitz.open(str(ruta))
    try:
        if doc.needs_pass:  # protegido con contraseña (ING-06)
            return ResultadoExtraccion(
                ruta=str(ruta), problema="El PDF está protegido con contraseña."
            )
        paginas_texto: list[str] = [doc[i].get_text("text") for i in range(doc.page_count)]
        n_paginas = doc.page_count
    finally:
        doc.close()

    texto_total = "\n".join(paginas_texto).strip()
    if not texto_total:
        # PDF escaneado sin texto extraíble (ING-06): equivale a ausente para EVA.
        return ResultadoExtraccion(
            ruta=str(ruta),
            problema="El PDF no tiene texto extraíble (¿escaneado sin OCR?).",
        )

    secciones = _detectar_secciones(paginas_texto)
    contenido = ContenidoDocumento(
        tipo_artefacto=tipo_artefacto,
        texto=texto_total,
        secciones=secciones,
        paginas=n_paginas,
        texto_primera_pagina=paginas_texto[0] if paginas_texto else "",
    )
    return ResultadoExtraccion(ruta=str(ruta), contenido=contenido)


def _detectar_secciones(paginas_texto: list[str]) -> tuple[Seccion, ...]:
    """Arma secciones a partir de líneas que parecen encabezados numerados."""
    secciones: list[Seccion] = []
    cuerpo_actual: list[str] = []

    def cerrar(sec: Seccion | None) -> None:
        if sec is not None:
            secciones[-1] = Seccion(
                numero=sec.numero, titulo=sec.titulo, pagina=sec.pagina,
                texto="\n".join(cuerpo_actual).strip(),
            )

    actual: Seccion | None = None
    for i, texto_pag in enumerate(paginas_texto, start=1):
        for linea in texto_pag.splitlines():
            m = _RE_ENCABEZADO.match(linea)
            if m and len(m.group(2).split()) <= 12:
                cerrar(actual)
                actual = Seccion(numero=m.group(1), titulo=m.group(2).strip(), pagina=i)
                secciones.append(actual)
                cuerpo_actual = []
            elif actual is not None:
                cuerpo_actual.append(linea)
    cerrar(actual)
    return tuple(secciones)


# ----------------------------------------------------------------------------
# docx
# ----------------------------------------------------------------------------
def _extraer_docx(ruta: Path, tipo_artefacto: str) -> ResultadoExtraccion:
    import docx  # python-docx

    documento = docx.Document(str(ruta))
    parrafos = [p.text for p in documento.paragraphs]
    texto_total = "\n".join(parrafos).strip()
    if not texto_total:
        return ResultadoExtraccion(
            ruta=str(ruta), problema="El .docx no tiene texto."
        )

    secciones: list[Seccion] = []
    cuerpo: list[str] = []
    for p in documento.paragraphs:
        es_encabezado = (p.style.name or "").lower().startswith("heading")
        m = _RE_ENCABEZADO.match(p.text)
        if es_encabezado or (m and len(m.group(2).split()) <= 12):
            if secciones:
                s = secciones[-1]
                secciones[-1] = Seccion(s.numero, s.titulo, s.pagina, "\n".join(cuerpo).strip())
            numero = m.group(1) if m else None
            titulo = m.group(2).strip() if m else p.text.strip()
            secciones.append(Seccion(numero=numero, titulo=titulo, pagina=None))  # docx: sin página
            cuerpo = []
        elif secciones:
            cuerpo.append(p.text)
    if secciones:
        s = secciones[-1]
        secciones[-1] = Seccion(s.numero, s.titulo, s.pagina, "\n".join(cuerpo).strip())

    contenido = ContenidoDocumento(
        tipo_artefacto=tipo_artefacto,
        texto=texto_total,
        secciones=tuple(secciones),
        paginas=None,  # ING/Etapa 0.2: docx no tiene páginas hasta renderizar
        texto_primera_pagina="\n".join(parrafos[:20]),
    )
    return ResultadoExtraccion(ruta=str(ruta), contenido=contenido)


# ----------------------------------------------------------------------------
# .ui (Qt Designer)
# ----------------------------------------------------------------------------
def _extraer_ui(ruta: Path, tipo_artefacto: str) -> ResultadoExtraccion:
    raiz_xml = ET.parse(str(ruta)).getroot()
    widget_raiz = raiz_xml.find("widget")
    nodo = _nodo_desde_elemento(widget_raiz) if widget_raiz is not None else None
    return ResultadoExtraccion(
        ruta=str(ruta), contenido=ArbolUI(tipo_artefacto=tipo_artefacto, raiz=nodo)
    )


def _nodo_desde_elemento(elem: ET.Element) -> NodoUI:
    """Construye recursivamente el árbol de objetos desde un <widget>/<layout>."""
    hijos = tuple(_nodo_desde_elemento(c) for c in _hijos_estructurales(elem))
    return NodoUI(clase=elem.get("class", ""), nombre=elem.get("name"), hijos=hijos)


def _hijos_estructurales(elem: ET.Element):
    """Widgets y layouts hijos directos, saltando los envoltorios <item>."""
    for c in elem:
        if c.tag in ("widget", "layout"):
            yield c
        elif c.tag == "item":
            for cc in c:
                if cc.tag in ("widget", "layout"):
                    yield cc
