"""Tests del pipeline analizar_entrega (Arquitectura §6, EVA-10).

El test estrella es el de la desconexión: cortar en una unidad, reanudar, y que
las unidades ya completadas no se vuelvan a pagar.
"""

from __future__ import annotations

import json

from qhaway.dominio.estados import EstadoEvaluacion as E
from qhaway.dominio.niveles import Nivel
from qhaway.dominio.rubrica import Rubrica
from qhaway.infra import crear_ciclo
from qhaway.infra.conector_ia import ConectorFalso, ErrorTransitorio
from qhaway.servicios import ContextoAnalisis, analizar_entrega

UI_XML = """<?xml version="1.0" encoding="UTF-8"?>
<ui version="4.0"><widget class="QWidget" name="Form">
 <widget class="QPushButton" name="guardar"/>
</widget></ui>
"""


def _rubrica():
    niveles = {n.value: "..." for n in Nivel}
    datos = {"rubrica": {"nombre": "R", "escala": {"tope_por_critico": 6}, "secciones": [
        {"artefacto": "presentacion", "criterios": [
            {"id": "PRE-1", "descripcion": "d", "peso": 1, "niveles": niveles}]},
        {"artefacto": "srs", "criterios": [
            {"id": "SRS-REQ", "descripcion": "d", "peso": 3, "niveles": niveles},
            {"id": "SRS-EST", "descripcion": "d", "peso": 1, "niveles": niveles}]},
        {"artefacto": "ui", "criterios": [
            {"id": "UI-NOM", "descripcion": "d", "peso": 2, "niveles": niveles}]},
    ]}}
    return Rubrica.desde_dict(datos, artefactos_requeridos=frozenset())


def _contexto():
    return ContextoAnalisis(
        rubrica=_rubrica(),
        checklists={},
        nomenclatura=__import__("qhaway.infra", fromlist=["NOMENCLATURA_DEFECTO"]).NOMENCLATURA_DEFECTO,
    )


def _preparar_entrega(tmp_path):
    """Ciclo + grupo + entrega con srs (.docx) y ui (.ui) presentes; presentacion ausente."""
    import docx

    c = crear_ciclo(tmp_path / "AED2-2027", "AED II — 2027")
    gid = c.grupos.crear(c.ciclo_id, "G01", "Grupo", "Proyecto")
    grupo = c.grupos.obtener(gid)
    entrega = c.entregas.crear_version(gid, 1, "2027-04-10", E.ENTREGA_CARGADA.value)

    cv = c.carpeta_version(grupo, 1, 1)
    (cv / "entrega").mkdir(parents=True, exist_ok=True)

    # srs.docx
    d = docx.Document()
    d.add_paragraph("1 Requerimientos Funcionales", style="Heading 1")
    d.add_paragraph("RF-01 el sistema deberá guardar.")
    ruta_srs = cv / "entrega" / "srs.docx"
    d.save(str(ruta_srs))
    c.archivos.agregar(entrega.id, "srs", c.rutas.relativa(ruta_srs), "docx")

    # form.ui
    ruta_ui = cv / "entrega" / "form.ui"
    ruta_ui.write_text(UI_XML, encoding="utf-8")
    c.archivos.agregar(entrega.id, "ui", c.rutas.relativa(ruta_ui), "ui")

    # presentacion: sin archivo (ausente)
    return c, grupo, entrega


def _resp_srs():
    return json.dumps({"artefacto": "srs", "valoraciones": [
        {"criterio_id": "SRS-REQ", "nivel": "Bueno", "justificacion": "x"},
        {"criterio_id": "SRS-EST", "nivel": "Bueno", "justificacion": "x"}],
        "observaciones": [{"criterio_id": "SRS-REQ", "tipo": "mejora",
            "contenido": "mejorar RF-01", "referencia": {"ubicacion": "1 Requerimientos"}}]})


def _resp_ui():
    return json.dumps({"artefacto": "ui", "valoraciones": [
        {"criterio_id": "UI-NOM", "nivel": "Regular", "justificacion": "x"}],
        "observaciones": []})


def _resp_transversal():
    return json.dumps({"consistencias": [], "preguntas_defensa": [
        {"pregunta": "¿por qué guardar así?", "elemento": "guardar", "artefacto": "ui",
         "intencion": "comprensión"}], "senales": []})


def test_camino_feliz_completo(tmp_path):
    c, grupo, entrega = _preparar_entrega(tmp_path)
    conector = ConectorFalso([_resp_srs(), _resp_ui(), _resp_transversal()],
                             dormir=lambda s: None, reloj=lambda: "t")
    r = analizar_entrega(c, grupo, entrega, _contexto(), conector)

    assert r.estado_final == E.BORRADOR_EN_REVISION.value
    assert not r.unidades_pendientes
    assert set(r.unidades_completadas) == {"presentacion", "srs", "ui", "transversal"}
    assert r.nota is not None

    # presentacion ausente -> Insuficiente sin llamada (EVA-05)
    vals = c.valoraciones.de_evaluacion(r.evaluacion_id)
    assert vals["PRE-1"]["nivel_ia"] == Nivel.INSUFICIENTE.value
    # Se creó la pregunta de defensa transversal
    preguntas = c.elementos.de_evaluacion(r.evaluacion_id, "pregunta_defensa")
    assert len(preguntas) == 1
    # Solo hubo 3 llamadas (srs, ui, transversal); presentacion no llamó
    assert conector.llamadas == 3
    c.cerrar()


def test_artefacto_ausente_no_llama_api(tmp_path):
    c, grupo, entrega = _preparar_entrega(tmp_path)
    conector = ConectorFalso([_resp_srs(), _resp_ui(), _resp_transversal()],
                             dormir=lambda s: None, reloj=lambda: "t")
    r = analizar_entrega(c, grupo, entrega, _contexto(), conector)
    # La nota compone Insuficiente(2) para PRE-1 junto al resto
    assert r.nota is not None
    c.cerrar()


def test_desconexion_reanudacion_no_repaga(tmp_path):
    c, grupo, entrega = _preparar_entrega(tmp_path)

    # Corte: srs OK, ui cae la red (se agotan los reintentos -> pendiente).
    guion_corte = [_resp_srs()] + [ErrorTransitorio("sin internet")] * 4
    conector1 = ConectorFalso(guion_corte, dormir=lambda s: None, reloj=lambda: "t")
    r1 = analizar_entrega(c, grupo, entrega, _contexto(), conector1)

    assert r1.estado_final == E.ANALISIS_INTERRUMPIDO.value
    assert "ui" in r1.unidades_pendientes or "transversal" in r1.unidades_pendientes
    assert "srs" in r1.unidades_completadas

    # Consumo de srs tras el corte: una llamada exitosa.
    srs_analisis_id = [
        f["id"] for f in c.analisis.unidades(r1.evaluacion_id) if f["unidad"] == "srs"
    ][0]
    consumos_srs_antes = c.con.execute(
        "SELECT COUNT(*) AS n FROM consumo_api WHERE analisis_id = ?", (srs_analisis_id,)
    ).fetchone()["n"]
    assert consumos_srs_antes == 1

    # --- Reanudación con un conector nuevo (la app se reabrió) ---
    conector2 = ConectorFalso([_resp_ui(), _resp_transversal()],
                              dormir=lambda s: None, reloj=lambda: "t")
    r2 = analizar_entrega(c, grupo, entrega, _contexto(), conector2)

    assert r2.estado_final == E.BORRADOR_EN_REVISION.value
    assert not r2.unidades_pendientes
    # CLAVE: srs NO se volvió a llamar; solo ui + transversal.
    assert conector2.llamadas == 2
    # Y srs no acumuló consumo nuevo: sigue en 1 (no se repagó).
    consumos_srs_despues = c.con.execute(
        "SELECT COUNT(*) AS n FROM consumo_api WHERE analisis_id = ?", (srs_analisis_id,)
    ).fetchone()["n"]
    assert consumos_srs_despues == 1
    c.cerrar()


def test_transversal_con_criterios_de_trazabilidad(tmp_path):
    """Rúbrica con sección transversal valorable (TRZ): se valora y compone nota."""
    import json as _json
    from qhaway.infra import NOMENCLATURA_DEFECTO
    from qhaway.infra.conector_ia import ConectorFalso
    from qhaway.servicios import ContextoAnalisis, analizar_entrega

    c = crear_ciclo(tmp_path / "c", "ciclo")
    gid = c.grupos.crear(c.ciclo_id, "G01", "N", "P")
    grupo = c.grupos.obtener(gid)
    entrega = c.entregas.crear_version(gid, 1, "2027-04-10", E.ENTREGA_CARGADA.value)
    cv = c.carpeta_version(grupo, 1, 1)
    (cv / "entrega").mkdir(parents=True, exist_ok=True)
    (cv / "entrega" / "form.ui").write_text(UI_XML, encoding="utf-8")
    c.archivos.agregar(entrega.id, "ui", c.rutas.relativa(cv / "entrega" / "form.ui"), "ui")

    niveles = {n.value: "..." for n in Nivel}
    rubrica = Rubrica.desde_dict({"rubrica": {"nombre": "R", "escala": {"tope_por_critico": 6},
        "secciones": [
            {"artefacto": "ui", "criterios": [
                {"id": "UI-NOM", "descripcion": "d", "peso": 2, "niveles": niveles}]},
            {"artefacto": "transversal", "criterios": [
                {"id": "TRZ-DOM", "descripcion": "Trazabilidad", "peso": 1, "niveles": niveles}]},
        ]}}, artefactos_requeridos=frozenset())
    contexto = ContextoAnalisis(rubrica=rubrica, checklists={}, nomenclatura=NOMENCLATURA_DEFECTO)

    resp_ui = _json.dumps({"artefacto": "ui", "valoraciones": [
        {"criterio_id": "UI-NOM", "nivel": "Bueno", "justificacion": "x"}], "observaciones": []})
    # La transversal ahora incluye valoraciones de trazabilidad
    resp_trans = _json.dumps({
        "valoraciones": [{"criterio_id": "TRZ-DOM", "nivel": "Regular", "justificacion": "x"}],
        "consistencias": [], "senales": [],
        "preguntas_defensa": [{"pregunta": "¿?", "elemento": "btnOk", "artefacto": "ui", "intencion": "x"}]})
    conector = ConectorFalso([resp_ui, resp_trans], dormir=lambda s: None, reloj=lambda: "t")

    r = analizar_entrega(c, grupo, entrega, contexto, conector)
    assert r.estado_final == E.BORRADOR_EN_REVISION.value
    assert r.nota is not None
    # La valoración de trazabilidad quedó registrada
    vals = c.valoraciones.de_evaluacion(r.evaluacion_id)
    assert vals["TRZ-DOM"]["nivel_ia"] == Nivel.REGULAR.value
    assert vals["UI-NOM"]["nivel_ia"] == Nivel.BUENO.value
    c.cerrar()
