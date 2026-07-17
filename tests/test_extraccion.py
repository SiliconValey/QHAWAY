"""Tests de extracción (ING-02/04/06) sobre archivos generados al vuelo.

Se generan los tres formatos reales para no depender de material externo, y se
corre DET de punta a punta sobre lo extraído.
"""

from __future__ import annotations

import pytest

from qhaway.dominio.deteccion import verificar_nomenclatura, verificar_completitud
from qhaway.infra import CHECKLIST_DEFECTO, NOMENCLATURA_DEFECTO, extraer


UI_XML = """<?xml version="1.0" encoding="UTF-8"?>
<ui version="4.0">
 <class>Form</class>
 <widget class="QWidget" name="Form">
  <layout class="QVBoxLayout" name="verticalLayout">
   <item>
    <widget class="QLabel" name="lblTitulo"/>
   </item>
   <item>
    <widget class="QLineEdit" name="nombreUsuario"/>
   </item>
   <item>
    <widget class="QPushButton" name="btnGuardar"/>
   </item>
   <item>
    <widget class="QPushButton" name="cancelar"/>
   </item>
  </layout>
 </widget>
</ui>
"""


def test_extraer_ui_construye_arbol(tmp_path):
    ruta = tmp_path / "form.ui"
    ruta.write_text(UI_XML, encoding="utf-8")
    res = extraer(ruta, "ui")
    assert res.ok
    arbol = res.contenido
    clases = [n.clase for n in arbol.widgets()]
    assert "QPushButton" in clases and "QLineEdit" in clases


def test_ui_nomenclatura_end_to_end(tmp_path):
    ruta = tmp_path / "form.ui"
    ruta.write_text(UI_XML, encoding="utf-8")
    arbol = extraer(ruta, "ui").contenido
    hallazgos = verificar_nomenclatura(arbol, NOMENCLATURA_DEFECTO)
    # 'nombreUsuario' (QLineEdit, falta txt) y 'cancelar' (QPushButton, falta btn)
    nombres = {h.datos["nombre_actual"] for h in hallazgos}
    assert nombres == {"nombreUsuario", "cancelar"}


def test_extraer_docx(tmp_path):
    import docx

    d = docx.Document()
    d.add_heading("Documento", level=0)
    d.add_paragraph("Instituto Superior. Integrantes: Ana, Beto. Profesor: X.")
    d.add_paragraph("1 Introducción", style="Heading 1")
    d.add_paragraph("Este documento describe el sistema.")
    d.add_paragraph("3 Requerimientos Funcionales", style="Heading 1")
    d.add_paragraph("RF-01: el sistema deberá...")
    ruta = tmp_path / "srs.docx"
    d.save(str(ruta))

    res = extraer(ruta, "srs")
    assert res.ok
    doc = res.contenido
    assert doc.paginas is None  # docx no tiene páginas
    titulos = [s.titulo for s in doc.secciones]
    assert any("Requerimientos" in t for t in titulos)
    # numeradas detectadas
    assert any(s.numerada for s in doc.secciones)


def test_extraer_pdf_y_secciones(tmp_path):
    import fitz

    doc = fitz.open()
    p1 = doc.new_page()
    p1.insert_text((72, 72), "Instituto Superior\nIntegrantes: Ana, Beto\nProfesor: X")
    p2 = doc.new_page()
    p2.insert_text(
        (72, 72),
        "1 Introduccion\nEste sistema resuelve un problema.\n"
        "3 Requerimientos Funcionales\nRF-01 el sistema debera guardar.",
    )
    ruta = tmp_path / "srs.pdf"
    doc.save(str(ruta))
    doc.close()

    res = extraer(ruta, "srs")
    assert res.ok
    contenido = res.contenido
    assert contenido.paginas == 2
    numeros = [s.numero for s in contenido.secciones]
    assert "1" in numeros and "3" in numeros


def test_pdf_end_to_end_det(tmp_path):
    import fitz

    doc = fitz.open()
    pg = doc.new_page()
    pg.insert_text((72, 72), "1 Introduccion\nContenido.\n3 Requerimientos Funcionales\nRF-01.")
    ruta = tmp_path / "srs.pdf"
    doc.save(str(ruta))
    doc.close()

    contenido = extraer(ruta, "srs").contenido
    hallazgos = verificar_completitud(contenido, CHECKLIST_DEFECTO["srs"])
    tipos_ausentes = {h.datos["bloque"] for h in hallazgos}
    # Faltan varios bloques del checklist por defecto (p.ej. definiciones, referencias)
    assert "referencias" in tipos_ausentes
    assert "requerimientos_funcionales" not in tipos_ausentes  # ese sí está


def test_formato_no_soportado_ing02(tmp_path):
    ruta = tmp_path / "notas.txt"
    ruta.write_text("hola", encoding="utf-8")
    res = extraer(ruta, "srs")
    assert not res.ok
    assert "no soportado" in res.problema.lower()


def test_pdf_sin_texto_ing06(tmp_path):
    import fitz

    doc = fitz.open()
    doc.new_page()  # página en blanco, sin texto
    ruta = tmp_path / "escaneado.pdf"
    doc.save(str(ruta))
    doc.close()

    res = extraer(ruta, "srs")
    assert not res.ok
    assert "texto" in res.problema.lower()


def test_archivo_inexistente_no_aborta_ing06(tmp_path):
    res = extraer(tmp_path / "no-existe.pdf", "srs")
    assert not res.ok
    assert res.problema  # se reporta, no se levanta excepción
