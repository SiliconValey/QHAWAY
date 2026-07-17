"""Arnés CLI descartable de QHAWAY (Etapa 2).

NO es la interfaz del producto (esa es Qt, Etapa 7). Es una herramienta mínima
para ejercitar los servicios sin UI, y de paso material de clase sobre separación
dominio/infraestructura: acá se ve que el mismo dominio y la misma persistencia
funcionan sin una sola línea de Qt.

Uso:
    python3 qhaway_cli.py crear   <ruta>
    python3 qhaway_cli.py listar  <ruta>
    python3 qhaway_cli.py demo    <ruta>     # flujo completo + regla de oro

El subcomando `demo` arma un ciclo con un grupo y una evaluación (calculando la
nota con el dominio de la Etapa 1), materializa todo a archivos, borra la base y
la reconstruye desde las carpetas — la prueba viva del criterio de salida.
"""

from __future__ import annotations

import argparse
import shutil
import sys
from pathlib import Path

from qhaway.dominio import Nivel, Rubrica, calcular_nota
from qhaway.dominio.contenido import ArbolUI
from qhaway.dominio.deteccion import ejecutar_det_documento, ejecutar_det_ui
from qhaway.dominio.estados import EstadoEvaluacion as E
from qhaway.infra import (
    CHECKLIST_DEFECTO,
    NOMENCLATURA_DEFECTO,
    abrir_ciclo,
    crear_ciclo,
    extraer,
    reconstruir_desde_carpetas,
)
from qhaway.infra.config_usuario import cargar_clave, guardar_clave, probar_conexion


def _rubrica_demo() -> Rubrica:
    niveles = {n.value: "..." for n in Nivel}
    datos = {
        "rubrica": {
            "nombre": "AED II — Expo 1 (demo)",
            "escala": {"tope_por_critico": 6},
            "secciones": [
                {
                    "artefacto": "srs",
                    "criterios": [
                        {"id": "SRS-REQ", "descripcion": "Requisitos", "peso": 3,
                         "critico": True, "niveles": niveles},
                        {"id": "SRS-EST", "descripcion": "Estructura", "peso": 1,
                         "niveles": niveles},
                    ],
                },
            ],
        }
    }
    return Rubrica.desde_dict(datos, artefactos_requeridos=frozenset())


def cmd_crear(ruta: Path) -> None:
    c = crear_ciclo(ruta, "AED II — 2027")
    print(f"Ciclo creado en {ruta}  (base: {c.rutas.db.name})")
    c.cerrar()


def cmd_listar(ruta: Path) -> None:
    c = abrir_ciclo(ruta)
    grupos = c.grupos.listar(c.ciclo_id, incluir_archivados=True)
    if not grupos:
        print("Sin grupos.")
    for g in grupos:
        marca = " [archivado]" if g.archivado else ""
        print(f"  {g.codigo}  {g.nombre}  ({g.proyecto}){marca}")
    c.cerrar()


def cmd_demo(ruta: Path) -> None:
    if ruta.exists():
        shutil.rmtree(ruta)

    print("=== 1. Crear ciclo y grupo ===")
    c = crear_ciclo(ruta, "AED II — 2027")
    gid = c.grupos.crear(c.ciclo_id, "G03", "Los Andinos", "Distribuidora Andes")
    c.integrantes.agregar(gid, "Ana Pérez", "2027-03-01")
    c.integrantes.agregar(gid, "Beto Gómez", "2027-03-01")
    grupo = c.grupos.obtener(gid)
    print(f"  grupo {grupo.codigo} — {grupo.nombre} ({grupo.proyecto}), 2 integrantes")

    print("\n=== 2. Entrega + evaluación (nota vía dominio) ===")
    entrega = c.entregas.crear_version(gid, 1, "2027-04-10", E.BORRADOR_EN_REVISION.value)
    valoraciones = {"SRS-REQ": Nivel.BUENO, "SRS-EST": Nivel.REGULAR}
    comp = calcular_nota(_rubrica_demo(), valoraciones)
    print("  " + comp.explicacion().replace("\n", "\n  "))

    ev_id = c.evaluaciones.crear(entrega.id, E.BORRADOR_EN_REVISION.value)
    for crit, niv in valoraciones.items():
        c.valoraciones.registrar(ev_id, crit, niv.value, niv.value)
    c.evaluaciones.fijar_nota_sugerida(ev_id, comp.nota)
    c.evaluaciones.validar(ev_id, comp.nota, "2027-04-13")

    print("\n=== 3. Materializar todo a archivos ===")
    c.guardar_manifiesto_entrega(grupo, entrega, [
        {"nombre": "srs.pdf", "tipo_artefacto": "srs", "formato": "pdf"},
    ])
    cv = c.carpeta_version(grupo, 1, 1)
    (cv / "entrega" / "srs.pdf").write_bytes(b"%PDF-1.4 demo")
    c.archivos.agregar(entrega.id, "srs", c.rutas.relativa(cv / "entrega" / "srs.pdf"), "pdf")
    c.guardar_evaluacion_en_archivo(grupo, entrega, {
        "estado": E.EVALUACION_VALIDADA.value, "vigente": True, "fecha": "2027-04-10",
        "nota_sugerida": comp.nota, "nota_final": comp.nota,
        "fecha_validacion": "2027-04-13",
        "valoraciones": {
            k: {"nivel_ia": v.value, "nivel_final": v.value} for k, v in valoraciones.items()
        },
    })
    c.guardar_analisis_en_archivo(grupo, entrega, "srs", {"unidad": "srs", "observaciones": []})
    c.guardar_decisiones_en_archivo(grupo, entrega, {"decisiones": []})
    print(f"  integridad base↔carpetas: {'OK' if c.verificar_integridad().ok else 'FALLA'}")
    c.cerrar()

    print("\n=== 4. Cataclismo: borrar qhaway.db ===")
    (ruta / "qhaway.db").unlink()
    for wal in ruta.glob("qhaway.db-*"):
        wal.unlink()
    print("  base eliminada.")

    print("\n=== 5. Reconstruir desde carpetas (regla de oro) ===")
    rep = reconstruir_desde_carpetas(ruta)
    print("  " + rep.resumen().replace("\n", "\n  "))

    c2 = abrir_ciclo(ruta)
    ev = c2.con.execute("SELECT nota_final FROM evaluacion").fetchone()
    integ = c2.con.execute("SELECT COUNT(*) AS n FROM integrante").fetchone()["n"]
    print(f"\n  nota_final recuperada de archivo: {ev['nota_final']}")
    print(f"  integrantes recuperados: {integ}  (esperado 0: no viven en carpetas, AD-03)")
    c2.cerrar()


def cmd_det(ruta: Path, tipo: str) -> None:
    """Extrae un archivo y corre la verificación determinística (Capa 1, sin IA)."""
    res = extraer(ruta, tipo)
    if not res.ok:
        print(f"[problema] {res.problema}")  # ING-06 / ING-02
        return

    if isinstance(res.contenido, ArbolUI):
        hallazgos = ejecutar_det_ui(res.contenido, NOMENCLATURA_DEFECTO)
    else:
        checklist = CHECKLIST_DEFECTO.get(tipo)
        if checklist is None:
            print(f"No hay checklist por defecto para el tipo '{tipo}'.")
            return
        hallazgos = ejecutar_det_documento(res.contenido, checklist)

    print(f"=== DET [{tipo}] {ruta.name} — Capa 1 (determinística, sin IA) ===")
    if not hallazgos:
        print("  Sin hallazgos: cumple todo lo verificable determinísticamente.")
    for h in hallazgos:
        print(f"  [{h.categoria}/{h.tipo}] {h.detalle}")


def cmd_guardar_clave() -> None:
    """Guarda la clave de API en la config de usuario (CFG-07), sin echo."""
    import os
    from getpass import getpass

    clave = os.environ.get("ANTHROPIC_API_KEY") or getpass("Pegá tu clave de API (no se muestra): ")
    if not clave.strip():
        print("No se ingresó ninguna clave.")
        return
    ruta = guardar_clave(clave.strip())
    print(f"Clave guardada en {ruta} (fuera del repo, CFG-07).")


def cmd_probar_conexion() -> None:
    """Prueba la conexión con la API (CFG-08)."""
    if not cargar_clave():
        print("No hay clave configurada. Usá 'guardar-clave' primero.")
        return
    ok, mensaje = probar_conexion()
    print(("[OK] " if ok else "[FALLA] ") + mensaje)


def cmd_pipeline_demo(ruta: Path) -> None:
    """Corre el pipeline con el conector falso: camino feliz + desconexión."""
    import json
    import shutil

    from qhaway.dominio.niveles import Nivel
    from qhaway.dominio.rubrica import Rubrica
    from qhaway.infra import NOMENCLATURA_DEFECTO
    from qhaway.infra.conector_ia import ConectorFalso, ErrorTransitorio
    from qhaway.servicios import ContextoAnalisis, analizar_entrega

    if ruta.exists():
        shutil.rmtree(ruta)

    niveles = {n.value: "..." for n in Nivel}
    rubrica = Rubrica.desde_dict({"rubrica": {"nombre": "R", "escala": {"tope_por_critico": 6},
        "secciones": [
            {"artefacto": "srs", "criterios": [
                {"id": "SRS-REQ", "descripcion": "d", "peso": 3, "niveles": niveles}]},
            {"artefacto": "ui", "criterios": [
                {"id": "UI-NOM", "descripcion": "d", "peso": 2, "niveles": niveles}]},
        ]}}, artefactos_requeridos=frozenset())
    contexto = ContextoAnalisis(rubrica=rubrica, checklists={}, nomenclatura=NOMENCLATURA_DEFECTO)

    c = crear_ciclo(ruta, "AED II — 2027")
    gid = c.grupos.crear(c.ciclo_id, "G01", "Grupo", "Proyecto")
    grupo = c.grupos.obtener(gid)
    entrega = c.entregas.crear_version(gid, 1, "2027-04-10", "entrega_cargada")
    cv = c.carpeta_version(grupo, 1, 1)
    (cv / "entrega").mkdir(parents=True, exist_ok=True)
    (cv / "entrega" / "form.ui").write_text(
        '<ui version="4.0"><widget class="QWidget" name="F">'
        '<widget class="QPushButton" name="guardar"/></widget></ui>', encoding="utf-8")
    c.archivos.agregar(entrega.id, "ui", c.rutas.relativa(cv / "entrega" / "form.ui"), "ui")
    import docx
    d = docx.Document(); d.add_paragraph("1 Requerimientos", style="Heading 1")
    d.add_paragraph("RF-01."); d.save(str(cv / "entrega" / "srs.docx"))
    c.archivos.agregar(entrega.id, "srs", c.rutas.relativa(cv / "entrega" / "srs.docx"), "docx")

    srs = json.dumps({"artefacto": "srs", "valoraciones": [
        {"criterio_id": "SRS-REQ", "nivel": "Bueno", "justificacion": "x"}], "observaciones": []})
    ui = json.dumps({"artefacto": "ui", "valoraciones": [
        {"criterio_id": "UI-NOM", "nivel": "Regular", "justificacion": "x"}], "observaciones": []})
    trans = json.dumps({"consistencias": [], "senales": [], "preguntas_defensa": [
        {"pregunta": "¿por qué?", "elemento": "guardar", "artefacto": "ui", "intencion": "x"}]})

    def progreso(unidad, estado):
        print(f"    · {unidad}: {estado}")

    print("=== Corte de red en la unidad 'ui' ===")
    c1 = ConectorFalso([srs] + [ErrorTransitorio("sin internet")] * 4,
                       dormir=lambda s: None, reloj=lambda: "t")
    r1 = analizar_entrega(c, grupo, entrega, contexto, c1, on_progreso=progreso)
    print(f"  estado: {r1.estado_final} | pendientes: {r1.unidades_pendientes} | costo: {r1.costo_total:.5f}")

    print("\n=== Se reabre la app y se reanuda ===")
    c2 = ConectorFalso([ui, trans], dormir=lambda s: None, reloj=lambda: "t")
    r2 = analizar_entrega(c, grupo, entrega, contexto, c2, on_progreso=progreso)
    print(f"  estado: {r2.estado_final} | nota sugerida: {r2.nota} | costo total: {r2.costo_total:.5f}")
    print(f"  llamadas en la reanudación: {c2.llamadas} (srs NO se repagó)")
    c.cerrar()


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Arnés CLI de QHAWAY (Etapas 2-5).")
    sub = parser.add_subparsers(dest="cmd", required=True)
    for nombre in ("crear", "listar", "demo", "pipeline-demo"):
        p = sub.add_parser(nombre)
        p.add_argument("ruta", type=Path)
    p_det = sub.add_parser("det", help="Extraer + verificación determinística de un archivo")
    p_det.add_argument("ruta", type=Path)
    p_det.add_argument("--tipo", default="srs", help="presentacion|srs|fd|ui")
    sub.add_parser("guardar-clave", help="Guardar la clave de API (CFG-07)")
    sub.add_parser("probar-conexion", help="Probar la conexión con la API (CFG-08)")
    args = parser.parse_args(argv)

    if args.cmd == "det":
        cmd_det(args.ruta, args.tipo)
    elif args.cmd == "guardar-clave":
        cmd_guardar_clave()
    elif args.cmd == "probar-conexion":
        cmd_probar_conexion()
    elif args.cmd == "pipeline-demo":
        cmd_pipeline_demo(args.ruta)
    else:
        {"crear": cmd_crear, "listar": cmd_listar, "demo": cmd_demo}[args.cmd](args.ruta)
    return 0


if __name__ == "__main__":
    sys.exit(main())
